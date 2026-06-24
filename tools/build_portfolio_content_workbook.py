from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_PATH = OUT_DIR / "个人简历网站_内容与照片补充工作簿_已确认版.docx"

BLUE = "2563EB"
DARK = "172033"
MUTED = "667085"
LIGHT_BLUE = "EAF1FF"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "D0D5DD"
GREEN = "18794E"
AMBER = "A15C00"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_no_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None, italic=None, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in [
        ("Heading 1", 18, DARK, 16, 8),
        ("Heading 2", 13.5, BLUE, 12, 5),
        ("Heading 3", 11.5, DARK, 9, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Portfolio Content Workbook  •  Zhuofan Yan")
    set_run_font(run, size=8.5, color=MUTED)


def add_label(doc, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    set_run_font(run, size=8.5, bold=True, color=color)
    return p


def add_body(doc, text, bold_prefix=None, color=DARK):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)
    return p


def add_bullet(doc, text, color=DARK):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r, color=color)
    return p


def add_prompt(doc, prompt, hint=None, lines=2):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], 120)
    set_table_borders(table, color="B9C8E8", size="7")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("请补充｜" + prompt)
    set_run_font(r, size=10, bold=True, color=BLUE)
    if hint:
        hp = cell.add_paragraph()
        hp.paragraph_format.space_after = Pt(3)
        hr = hp.add_run(hint)
        set_run_font(hr, size=9, italic=True, color=MUTED)
    for _ in range(lines):
        lp = cell.add_paragraph("________________________________________________________________________________")
        lp.paragraph_format.space_after = Pt(2)
        for run in lp.runs:
            set_run_font(run, size=9, color="AEB7C4")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_photo_request(doc, photo_id, purpose, guidance, optional=False):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], 120)
    set_table_borders(table, color="F0B36A", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7E8")
    p = cell.paragraphs[0]
    r = p.add_run(f"需要上传照片｜{photo_id}  {'（可选）' if optional else '（必需）'}")
    set_run_font(r, size=10, bold=True, color=AMBER)
    p2 = cell.add_paragraph()
    r2 = p2.add_run("用途：" + purpose)
    set_run_font(r2, size=9.5, bold=True, color=DARK)
    p3 = cell.add_paragraph()
    r3 = p3.add_run("建议：" + guidance)
    set_run_font(r3, size=9, color=MUTED)
    p4 = cell.add_paragraph()
    r4 = p4.add_run("文件名：________________________________    拍摄时间/地点：____________________________")
    set_run_font(r4, size=9, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_status_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], 120)
    set_table_borders(table, color="A8D5C1", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "ECF8F2")
    p = cell.paragraphs[0]
    r = p.add_run("已有信息｜" + text)
    set_run_font(r, size=9.5, color=GREEN)


def add_prefill(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], 120)
    set_table_borders(table, color="98A2B3", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    p = cell.paragraphs[0]
    r = p.add_run("已从简历预填｜")
    set_run_font(r, size=9.5, bold=True, color=MUTED)
    r2 = p.add_run(text)
    set_run_font(r2, size=9.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_verify(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], 120)
    set_table_borders(table, color="E6A23C", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7E8")
    p = cell.paragraphs[0]
    r = p.add_run("需要你核实｜" + text)
    set_run_font(r, size=9.5, bold=True, color=AMBER)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_section_intro(doc, number, title, purpose):
    add_label(doc, f"SECTION {number}")
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(purpose)
    set_run_font(r, size=11, italic=True, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(86)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PORTFOLIO WEBSITE")
    set_run_font(r, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("个人简历网站\n内容与照片补充工作簿")
    set_run_font(r, size=28, bold=True, color=DARK, name="Aptos Display")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Zhuofan Yan  ·  Website Content Collection")
    set_run_font(r, size=12, color=MUTED)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [7600], 1000)
    set_table_borders(table, color="B9C8E8", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "使用方式\n"
        "从第 1 部分开始逐项填写；遇到橙色框时上传对应照片。\n"
        "不知道最终英文怎么写也没关系，先用中文讲清楚事实、故事和你的判断。"
    )
    set_run_font(r, size=11, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("版本：2026-06-22")
    set_run_font(r, size=9, color=MUTED)


def add_quick_start(doc):
    add_section_intro(doc, "00", "填写说明与优先级", "这份工作簿的目标，是收齐网站所需的事实、故事、数据和视觉素材。")
    doc.add_heading("建议填写顺序", level=2)
    for item in [
        "第一轮：Hero、About Me、Contact——30 分钟内完成网站最基础的信息。",
        "第二轮：Experience 六页——每页先写一个真实瞬间，不追求漂亮文案。",
        "第三轮：Cases——先核实数字和个人贡献，再补完整过程。",
        "第四轮：Me & AI——提供网站链接、截图和真实工作流。",
        "最后：按文末“照片上传清单”统一整理文件名并上传。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("填写原则", level=2)
    for item in [
        "事实优先：时间、地点、角色、个人贡献和结果必须准确。",
        "具体优先：用一个场景或决定代替泛泛的“提升了能力”。",
        "个人贡献优先：团队成果旁边必须说清楚你具体做了什么。",
        "反思优先：案例中保留 How I’d do it better now，体现成长与判断。",
        "可公开优先：公司内部材料、数据和截图请先确认是否能公开。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("网站叙事主线", level=2)
    add_body(
        doc,
        "我如何观察人、理解需求，把洞察变成产品，并让产品真正进入市场。",
    )
    add_prompt(
        doc,
        "你希望招聘者看完网站后，用哪一句话描述你？",
        "例如：她不只会做营销，也能把用户洞察转化为真正落地的产品。",
        lines=3,
    )


def add_hero(doc):
    add_section_intro(doc, "01", "Hero / Intro", "目标：让访客在 10 秒内知道你是谁、你创造什么价值，以及为什么值得继续看。")
    add_status_callout(
        doc,
        "网站显示姓名已填写为 Jovan Yan；主文案方向已确定：I turn observations about people into products, stories, and experiences they want to join."
    )
    add_prefill(
        doc,
        "简历姓名为“阎卓凡”。当前职业概况：2 年快消产品管理经验，擅长将消费者洞察与市场趋势转化为产品概念，并推进端到端 GTM；同时具备内容营销、达人种草、视觉物料把控和跨国协同经验。"
    )
    add_prefill(
        doc,
        "目标岗位优先级（按最新填写同步）：1）Product Marketing；2）Product Manager；3）GTM。网站以 Product Marketing 为主要职业入口，同时强调产品定义与落地能力，避免被理解为只做传播执行。"
    )
    add_prefill(
        doc,
        "行业偏好：不限定特定行业。求职时更关注问题是否值得解决、团队是否重视用户、岗位是否能覆盖从洞察到落地的完整过程。网站不会把职业身份限定为快消或宠物行业，而会将相关经历作为可迁移的产品判断、跨团队协作和 GTM 能力证据。"
    )
    add_prefill(
        doc,
        "地点偏好：工作城市和地区不是主要筛选条件，对本地、异地、远程及跨区域机会保持开放。网站仅展示当前所在地作为基础信息，不将地域写成职业定位的一部分。"
    )
    add_prefill(
        doc,
        "Hero 职业关键词：Product Marketing、Community、Consumer Insight、GTM。"
    )
    doc.add_heading("建议显示文案", level=2)
    add_label(doc, "Primary statement")
    add_body(doc, "I turn observations about people into products, stories, and experiences they want to join.")
    add_label(doc, "Supporting line")
    add_body(doc, "I observe, research, build, and bring ideas to life.")
    add_prompt(
        doc,
        "Hero 辅助信息",
        "职业关键词已经填写。请继续确认是否展示 Open to opportunities，以及按钮文字。",
        lines=2,
    )
    add_prompt(
        doc,
        "主文案是否需要修改",
        "若保留可写“保留”；若修改，请直接写新的英文或先写中文意思。",
        lines=2,
    )
    add_photo_request(
        doc,
        "HERO-01",
        "首页主视觉个人照片",
        "竖版或接近 4:5；人物清晰、背景不过度杂乱；建议原图宽度至少 1600px。请提供 2–3 张备选，避免自拍滤镜过重。",
    )


def add_about(doc):
    add_section_intro(doc, "02", "About Me", "目标：解释你为何同时具备 Product 与 GTM 视角，并自然引出后面的经历故事。")
    doc.add_heading("建议章节标题", level=2)
    add_body(doc, "I work where product thinking meets go-to-market.")
    doc.add_heading("建议内容结构", level=2)
    for item in [
        "我是谁：兼具产品和营销背景的 Product / GTM Builder。",
        "我擅长什么：用户研究、洞察转译、产品落地、内容与市场进入。",
        "我的不同之处：既会思考人为什么行动，也理解产品怎样在现实约束中被做出来。",
        "故事引子：职业路径并非直线，而是从传播、研究和市场逐渐走向产品。",
    ]:
        add_bullet(doc, item)
    add_prefill(
        doc,
        "可用故事主线：从快手的内容传播，到百事/桂格的新品孵化与整合营销，再到 ZURU 的消费者洞察、产品定义、成本控制和全球新品交付；之后用 AI 辅助从零搭建内容获客流程。这条路径能够证明你并非只做营销，而是逐步形成 Product × GTM 的完整闭环。"
    )
    add_prefill(
        doc,
        "现有能力证据：用户与竞品研究；产品概念和商业可行性评估；外观、包材与陈列把控；设计/供应链/工厂跨团队协作；成本与销量测算；小红书、抖音、微博达人与内容策略；AI 辅助研究和内容生产。"
    )
    add_prefill(
        doc,
        "职业定位：优先应聘 Product Marketing，其次是 Product Manager 与 GTM。核心优势不是分别做过三类工作，而是能够把用户和市场信号转化为产品判断，并继续推动产品被正确表达、上市和采用。"
    )
    add_prefill(
        doc,
        "希望避免的误解：我不是一个只会写方案、停留在策略表达层面的候选人。我的工作方式是从研究和洞察开始，把判断推进到产品定义、成本与可行性评估、设计打样、供应链与工厂协作、上市执行和结果复盘；方案对我而言是推动落地的工具，而不是工作的终点。"
    )
    add_prefill(
        doc,
        "职业身份补充：我同时理解产品与营销。营销和用户洞察背景让我在产品定位阶段能够更深入地理解需求，并提前思考分发与转化路径；产品开发经验则让我在做营销时更清楚产品真正的功能、优势与限制，能够把卖点与消费者的具体需求连接起来。"
    )
    add_prefill(
        doc,
        "PM × GTM 的价值：我通过定性和定量研究寻找用户已经表达和尚未说出口的需求。在 ZURU，我把年轻消费者对宠物的情感依赖理解为对情绪价值、拟人化和互动体验的需求，因此在产品设计中参考儿童产品的趣味性与互动方式；在百事/桂格，我结合对配方和产品开发的理解，把营养与成分优势拆解到不同人群和场景中，让营销从泛化传播变成对具体痛点的回应。"
    )
    add_prefill(
        doc,
        "已确认的能力案例 01：从线下手账活动的反馈中识别“五年日记难坚持”的问题，在便携、可视化反馈、关键词检索与手写仪式感之间做产品取舍，并独立构建“时间锚点”原型。"
    )
    add_prefill(
        doc,
        "已确认的能力案例 02：把一个 idea 从 0 推进到 1。在 ZURU，我能够把初步产品想法拆解为消费者洞察、产品定义、成本与商业可行性、外观与包材、打样试产和交付节点，并协调设计、供应链、工厂与原料商解决现实约束，让概念最终成为可生产、可销售的真实产品。"
    )
    add_prefill(
        doc,
        "已确认的能力案例 03：把复杂产品转化为用户能够理解的价值。在百事/桂格项目中，我不仅了解产品的营养、成分和配方，也把这些专业卖点拆解到不同人群与使用场景中，再转化为用户能迅速理解的内容主题、达人话术和传播表达，让产品优势真正对应消费者的日常痛点与需求。"
    )
    add_prompt(
        doc,
        "你的故事主线（100–200 字中文）",
        "从本科开始，你是如何一步步走到今天的？这里只要主线，不需要列完整履历。",
        lines=6,
    )
    add_photo_request(
        doc,
        "ABOUT-01",
        "About Me 辅助照片",
        "可选工作状态、观察城市、写作或与人交流的自然抓拍；横版优先。不要与 Hero 使用同一张照片。",
        optional=True,
    )


def add_experience_page(doc, page_no, title, known, photo_id, prompts, prefill=None, verify=None):
    doc.add_heading(f"Page {page_no} — {title}", level=2)
    if known:
        add_status_callout(doc, known)
    if prefill:
        add_prefill(doc, prefill)
    if verify:
        add_verify(doc, verify)
    add_prompt(
        doc,
        "这一页的基础信息",
        "地点｜起止时间（月/年）｜学校/公司/组织｜你的身份或职位。",
        lines=2,
    )
    add_prompt(
        doc,
        "这一页要讲的一个故事",
        "请写一个具体瞬间：发生了什么、你当时观察到什么、你做了什么。建议 100–200 字中文。",
        lines=5,
    )
    add_prompt(
        doc,
        "What I carried forward",
        "这段经历留下了什么能力、判断或工作习惯？它如何连接到下一阶段？",
        lines=3,
    )
    for prompt, hint in prompts:
        add_prompt(doc, prompt, hint, lines=2)
    add_photo_request(
        doc,
        photo_id,
        f"Experience Page {page_no} 主照片",
        "优先选择能代表当时环境和故事的真实照片，而不是公司 logo。横版 3:2 或 4:3 最灵活；原图宽度建议至少 1600px。",
    )


def add_experience(doc):
    add_section_intro(doc, "03", "Experience — My Journey Book", "目标：用一本六页的“书”讲职业形成过程；每一页只承担一个核心转折。")
    add_status_callout(doc, "最终结构已确认：目录、本科、快手实习、百事实习、研究生、ZURU 工作，共六页。")
    doc.add_heading("整本书的统一规则", level=2)
    for item in [
        "每页：1 张照片 + 地点 + 时间 + 1 个短故事 + 1 个成长点。",
        "故事不要写成岗位职责；优先写一次观察、一次决定或一次转变。",
        "每页默认可见文字建议控制在 45–80 个英文单词。",
        "公司成果可提及，但完整项目拆解留到 Cases。",
    ]:
        add_bullet(doc, item)
    add_prompt(
        doc,
        "这本书的总标题",
        "暂定 Notes From the Places That Shaped Me。可保留、修改，或先写你希望表达的中文意思。",
        lines=2,
    )
    add_experience_page(
        doc, "01", "Contents / 目录",
        "目录页用于展示六个章节和路线，不承担完整故事。",
        "EXP-01",
        [
            ("六页的最终章节名称", "请按 1–6 顺序列出，希望更像故事章节，而不是简历条目。"),
        ],
        prefill="已确认节点：华东师范大学（2018.09–2022.06）→ 北京快手（2021.02–2021.06）→ 百事中国（2022.03–2022.06）→ 波士顿大学（2022.09–2024.01）→ ZURU（2024.07–2026.01）。城市路线：Shanghai → Beijing → Boston → Shenzhen。",
    )
    add_experience_page(
        doc, "02", "本科 / Where curiosity began",
        "已知线索：East China Normal University、Business Administration、AIESEC、校园组织与早期社区经历，均需你核实。",
        "EXP-02",
        [
            ("本科最重要的经历", "AIESEC、社团、社区或课程中，哪件事最能解释你后来为什么关注“参与感”？"),
        ],
        prefill="East China Normal University｜Business Administration 本科｜上海，中国｜2018 年 9 月–2022 年 6 月。AIESEC 华东师范大学支持团队海外志愿者市场部负责人：策划线上活动，独立与日本学校沟通；运营 90 人微信招新群，提升群内活跃度并超额完成招新目标。",
        verify="AIESEC 经历已补充。仍需说明这段经历让你形成了什么关于用户参与、社区运营或跨文化沟通的判断。",
    )
    add_experience_page(
        doc, "03", "实习一 / First contact with real users",
        "建议使用第一段实习：北京快手科技有限公司｜公关传播实习生。",
        "EXP-03",
        [
            ("公司、部门与职位", "请写官方英文名称，并说明这是否是你的第一段实习。"),
            ("为什么把它放在第一章", "它第一次让你理解了用户、内容、市场或商业中的什么？"),
        ],
        prefill="北京｜2021 年 2 月–2021 年 6 月。参与平台内容传播项目；独立完成热点选题、内容方案、达人资源引入与数据跟进。女性主题内容带来微博话题阅读量 1 亿+、账号单日涨粉 10.7 万；参与丁真直播策划，观看 280 万+、点赞 87 万+、转粉率 11%；筛选并引入 20+ 创作者。",
        verify="请补充一个具体故事：你如何捕捉女性用户情绪，或在丁真直播策划中提出过什么具体想法。团队总数据旁还需说明你的个人贡献。",
    )
    add_experience_page(
        doc, "04", "实习二 / Learning how ideas travel",
        "建议使用第二段实习：百事中国｜产品营销实习生。",
        "EXP-04",
        [
            ("公司、部门与职位", "请写官方英文名称、实习时间和工作城市。"),
            ("一个可讲的成果或事件", "不要只写曝光量；说明你从受众反应中学到了什么。"),
        ],
        prefill="上海｜2022 年 3 月–2022 年 6 月。参与桂格健康快消新品孵化与 GTM，覆盖概念、卖点、上市筹备与上市后优化；制定小红书、抖音、微博整合营销方案，搭建达人矩阵。简历记录：4500 万+ 曝光、CPM 22–30、爆文率 24%；天猫小程序上线当天访问 5000+，两周内新品预约人数 20 万+。",
        verify="已确认使用 4500 万+ 曝光。仍需确认“主导新品端到端开发”在实习岗位中的准确职责边界。",
    )
    add_experience_page(
        doc, "05", "研究生 / Connecting research with product thinking",
        "已知线索：Boston University、Marketing Research、跨文化生活与消费者研究方法，均需你核实。",
        "EXP-05",
        [
            ("研究生信息核实", "学校、项目/专业官方英文名、入学与毕业时间。"),
            ("研究如何改变你", "请选一个课程、研究项目或生活经历，说明它如何改变你的提问和判断方式。"),
        ],
        prefill="波士顿大学｜市场营销调研硕士｜波士顿，美国｜2022 年 9 月–2024 年 1 月。",
        verify="这一页希望重点介绍“在学校学了什么”。仍需补充具体课程、研究方法或项目，以及它如何改变了你理解消费者和验证判断的方法。",
    )
    add_experience_page(
        doc, "06", "工作 / Building at the intersection of Product & GTM",
        "已知线索：ZURU、Walmart、consumer research、concept、sampling、production、launch；需要核实职位与范围。",
        "EXP-06",
        [
            ("工作信息核实", "公司、职位、时间、地点、负责市场与产品类别。"),
            ("最能代表你的职责", "从用户研究到产品上市，你真正拥有或推动的是哪一段？"),
            ("下一步职业方向", "你希望进入什么行业、团队和岗位？接下来想解决什么问题？"),
        ],
        prefill="ZURU｜NPD Junior Project Management（新产品项目经理）｜深圳｜2024 年 7 月–2026 年 1 月。负责消费者洞察、产品定义、视觉与包材、商业可行性、设计/供应链/海外工厂协同、打样试产、成本和上市反馈；独立承接沃尔玛定制猫粮罐头产品线，并管理 2 名实习生的数据录入与问卷汇总。",
        verify="已确认 140 万为订单量。仍需补充：这些订单对应的产品、市场、客户范围和统计时间。",
    )


def add_case_template(doc, case_no, title, known, photo_prefix, prefill=None, verify=None):
    doc.add_heading(f"Case {case_no} — {title}", level=2)
    add_status_callout(doc, known)
    if prefill:
        add_prefill(doc, prefill)
    if verify:
        add_verify(doc, verify)
    fields = [
        ("项目基础信息", "项目时间｜公司/客户｜目标市场｜产品或活动｜我的角色。", 2),
        ("Context", "当时发生了什么？为什么这个项目重要？", 3),
        ("Challenge", "真正需要解决的问题是什么？不要只写任务说明。", 3),
        ("Insight", "你通过什么研究或观察发现了什么？请提供证据。", 4),
        ("Strategy & key decision", "你做出了什么关键选择？为什么没有选其他方案？", 4),
        ("Execution", "你本人具体做了什么？请区分个人贡献和团队贡献。", 5),
        ("Impact", "结果数据、口径、时间范围，以及是否可公开。", 3),
        ("Reflection", "你从项目中学到了什么？", 3),
        ("How I’d do it better now", "如果今天重做，你会在哪一步改变方法？为什么？", 4),
    ]
    for prompt, hint, lines in fields:
        add_prompt(doc, prompt, hint, lines=lines)
    add_photo_request(
        doc,
        f"{photo_prefix}-01",
        f"Case {case_no} 卡片封面图",
        "优先成品、真实使用场景或最能代表结果的图；横版 16:10 或 3:2。",
    )
    add_photo_request(
        doc,
        f"{photo_prefix}-02",
        f"Case {case_no} 详情证据图",
        "请上传 3–6 张：研究材料、过程、产品/活动、数据结果。涉及公司内部信息时先脱敏并确认可公开。",
    )


def add_cases(doc):
    add_section_intro(doc, "04", "Selected Cases", "目标：先用数字快速建立可信度，再通过展开页面展示你的判断、执行和复盘能力。")
    doc.add_heading("首页数字快照", level=2)
    add_status_callout(
        doc,
        "已确认/候选数据：快手女性主题微博话题阅读量 1 亿+；百事/桂格曝光 4500 万+；ZURU 沃尔玛猫粮罐头产品线首批订单 140 万；成本下降 3%–5%。"
    )
    add_verify(
        doc,
        "数字本身已按你的确认修正；仍需补充每个数字的统计时间、市场范围、团队成果与个人贡献。"
    )
    add_prompt(
        doc,
        "最终展示的 3–5 个数字",
        "每个数字请写：准确数值｜单位｜时间范围｜对应项目｜你的贡献｜是否可公开。",
        lines=6,
    )
    add_prompt(
        doc,
        "最终保留几个正式案例",
        "建议 3 个主案例。请写最终项目名称和优先顺序。",
        lines=3,
    )
    add_case_template(
        doc,
        "01",
        "ZURU × Walmart",
        "已确认：沃尔玛定制猫粮罐头产品线面向欧美市场，2025 年确认了对应 2026 年第一季度交付的首批订单（first order），数量为 140 万；综合成本下降 3%–5%；小组基础业务处理效率提升 20%。你是该产品线的端到端负责人。",
        "CASE01",
        prefill="沃尔玛定制猫粮罐头产品线：面向欧美市场。你个人端到端负责产品前期调研、用户与市场洞察、产品定位、概念创新与开发，并持续推进设计、打样、试产和正式生产；产品落地后，你还负责后期线下宣传。设计、供应链、工厂和原料商等专业执行由你协调相关团队共同完成。原料替换使综合成本下降 3%–5%，统一归档规范使小组基础业务效率提升 20%。该产品线于 2025 年确认了对应 2026 年第一季度交付的首批订单（first order），数量为 140 万。上市后的销量未达到预期。复盘来看，产品虽然在外观上具有一定差异，但在猫罐头品类中仍不足以与其他品牌拉开距离；与此同时，低价格与较高质量的价值定位又限制了高成本配方研发和大规模适口性测试。How I’d do it better now：如果重新开始，我不会直接假设“更复杂的配方”就是答案，而会在立项前用概念测试、用户访谈和小规模货架/传播测试，比较消费者对包装、份量、使用场景、功能利益和配方差异的真实反应，先找到最能影响购买且成本可承受的差异化支点，再决定研发投入。",
        verify="本案例聚焦“沃尔玛定制猫粮罐头产品线”。市场、订单口径和个人职责范围均已确认。",
    )
    add_case_template(
        doc,
        "02",
        "PepsiCo / Quaker",
        "已确认使用：4500 万+ 曝光、爆文率 24%、两周内新品预约人数 20 万+。数据口径和个人贡献待补充。",
        "CASE02",
        prefill="桂格健康快消新品：目标人群为 20–35 岁年轻都市女性。参与概念构思、卖点提炼、上市筹备及上市后优化；制定小红书、抖音、微博整合营销方案，研究竞品种草策略和内容话术，并把控内容调性。关键决策有两项：第一，把营养、成分等专业卖点转译成目标人群能够代入的具体生活场景；第二，搭建不同层级的达人矩阵，让不同类型的创作者分别承担扩大认知、建立信任与推动转化的任务。How I’d do it better now：第一，把线上内容种草与线下试吃、体验或活动结合，让消费者不只理解卖点，也能实际感受产品；第二，在产品开发阶段更早、更深入地验证消费者需求，减少产品定义与真实购买动机之间的偏差；第三，不把传播局限在社交媒体种草，而是根据产品特点组合直播、信息流广告、创意短视频等更多触点。高蛋白燕麦脆本身不是一个天然容易激发兴趣的品类，因此需要更丰富的触点和表达方式反复建立场景、价值与购买理由。",
        verify="曝光已确认使用 4500 万+；20 万+ 已确认为两周内新品预约人数；产品正式名称不公开；目标人群已确认为 20–35 岁年轻都市女性；关键决策已确认为场景化营销与多层级达人矩阵。",
    )
    doc.add_heading("候选 Case 04 — Journaling Community", level=2)
    add_prompt(
        doc,
        "是否把手帐社区作为正式案例？",
        "若保留，请补充项目周期、社区规模、活动形式、你的角色、成果或用户反馈；若不保留，可并入 Journey 或 Me & AI。",
        lines=5,
    )
    add_photo_request(
        doc,
        "CASE04-01",
        "手帐社区项目素材",
        "活动现场、用户作品、互动机制、反馈截图等，建议 3–6 张。",
        optional=True,
    )


def add_ai(doc):
    add_section_intro(doc, "05", "Me & AI", "目标：证明 AI 是你研究、构建和学习的能力放大器，而不是装饰性关键词。")
    doc.add_heading("建议开场句", level=2)
    add_body(doc, "AI doesn’t replace my judgment. It shortens the distance between an observation and something people can experience.")
    add_prefill(
        doc,
        "简历中的真实 AI 使用：用 Codex 抓取行业资讯和竞品参数；人工整理竞品参数与宠物饲养评价后，用 GPT 提炼用户核心诉求与市场差异化痛点，并整理为标准化调研报表支撑产品迭代。"
    )
    add_prompt(
        doc,
        "你如何定义自己与 AI 的关系",
        "请用一个真实例子说明：AI 加速了什么，而哪些判断仍必须由你完成？",
        lines=4,
    )
    doc.add_heading("Build 01 — Five-Year Diary / Time Anchor", level=2)
    add_status_callout(doc, "现有定位：A five-year diary for meeting your past self.")
    add_prefill(
        doc,
        "项目来源：/Users/zhuofan/Desktop/5year-diary。产品中文名“时间锚点”，是一款围绕“五年同日回望”设计的本地优先日记原型。核心用户是愿意持续记录、重视私密性，并希望多年后回看同一天变化的个人用户。"
    )
    add_prefill(
        doc,
        "核心机制：用户每天写下一条文字记录，也可添加图片；同一日期按 2026–2030 五年时间轴排列，未来年份暂时锁定。产品支持昨日/明日切换、年度记忆图谱、历史补记、后续便利贴、全文搜索与往日回音，并提供本地数据导入/导出构想。"
    )
    add_prefill(
        doc,
        "产品与品牌设计：使用纸张、长河、轻舟和锚点形成统一隐喻；首页通过“百舸记忆图谱”展示全年记录密度；主体验把同一天的五年记录放在一条时间轴上，让普通日记从“保存今天”变成“与过去的自己重新相遇”。"
    )
    add_prefill(
        doc,
        "项目起点：在线下手账活动中，我听到多位参与者反馈，实体五年日记很难长期坚持。一方面，日记本体积较大、不便随身携带；另一方面，记录过程缺少及时反馈，使用者很难直观看见自己已经写了多久、积累了多少内容。于是我开始思考：能否把五年日记做成一个更容易持续使用的线上产品？"
    )
    add_prefill(
        doc,
        "产品回应：我加入年度记录图谱，让用户直观看到自己的记录轨迹；设计关键词搜索，让用户可以找出所有包含某个词的日子，重新发现多年生活中的主题和变化。数字化并不意味着抹掉手写的质感——主记录不能被无限反复修改，以保留落笔时的认真和时间痕迹；当用户产生新的理解时，可以通过“后续补记”继续回应过去的自己。产品试图在手写日记的仪式感与数字工具的便携、检索和反馈之间取得平衡。"
    )
    add_prefill(
        doc,
        "我的角色与过程：这是一个由我端到端独立完成的产品项目。我先在线下手账活动中观察用户并收集反馈，从“难携带、缺少即时反馈、难以坚持”中定义问题；随后完成产品定位与核心机制设计，包括五年同日时间轴、年度记录图谱、关键词搜索、主记录锁定和后续补记。之后由我继续完成信息架构、交互、视觉语言、品牌隐喻与界面文案，并借助 AI 将想法开发成可运行的前端原型，最后通过实际操作与产品审查发现问题、整理优先级并规划下一轮迭代。"
    )
    add_prefill(
        doc,
        "AI 协作方式：我使用 ChatGPT 协助梳理和核实产品需求、搭建功能与前端内容框架，并帮助我把线下观察转化为更清晰的产品结构；随后使用 Gemini 完成原型探索、前端视觉设计和部分代码实现。AI 缩短了从想法到可交互原型的距离，但需求优先级、功能取舍、体验原则、视觉方向以及最终验收仍由我判断和决定。"
    )
    add_prefill(
        doc,
        "用户验证状态：当前原型尚未进行外部用户测试，因此不展示虚构的使用数据或反馈。下一步计划邀请最初接触到的手账用户进行小规模可用性测试，重点验证首次使用时能否理解“五年同日回望”、记录与保存是否顺畅、年度图谱和关键词搜索是否真正增强坚持感，以及“主记录锁定 + 后续补记”是否符合手写日记用户的心理预期。"
    )
    add_prefill(
        doc,
        "当前状态与反思：这是一个可运行的前端原型，而非已上线成熟产品。审查发现，下一步应优先解决保存可靠性、备份恢复、隐私文案、闰日、动态五年周期、移动端日期导航、图片存储容量和无障碍问题，再扩展月度回顾、去年今日与五年册导出。"
    )
    for prompt, hint, lines in [
        ("公开访问链接", "当前只有本地项目路径。若有部署链接请提供；若没有，网站制作阶段再部署。", 2),
    ]:
        add_prompt(doc, prompt, hint, lines=lines)
    add_photo_request(
        doc,
        "AI-01",
        "五年日记网站主视觉",
        "首页或最能说明核心体验的清晰截图；桌面端与移动端各 1 张。",
    )
    image_items = [
        (OUT_DIR / "5year-diary-screenshots" / "01-cover.png", "项目截图 01｜时间锚点首页与年度记忆图谱"),
        (OUT_DIR / "5year-diary-screenshots" / "02-diary-view.png", "项目截图 02｜同一日期的五年时间轴与当日记录"),
    ]
    if all(image_path.exists() for image_path, _ in image_items):
        gallery = doc.add_table(rows=1, cols=2)
        gallery.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_geometry(gallery, [4680, 4680], 120)
        set_table_no_borders(gallery)
        for cell, (image_path, caption) in zip(gallery.rows[0].cells, image_items):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run()
            run.add_picture(str(image_path), width=Inches(3.0))
            cp = cell.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(caption)
            set_run_font(cr, size=8, color=MUTED, italic=True)
    add_photo_request(
        doc,
        "AI-02",
        "五年日记完整 Case 过程",
        "用户流、早期草图、迭代版本、测试反馈或代码/AI 协作过程，建议 4–8 张。",
    )
    doc.add_heading("How I leverage AI", level=2)
    add_body(doc, "Research → Synthesis → Ideation → Prototyping → Testing → Storytelling")
    for stage in ["Research", "Synthesis", "Ideation", "Prototyping", "Testing", "Storytelling"]:
        add_prompt(
            doc,
            f"{stage}：一个真实工作流例子",
            "输入是什么｜AI 做了什么｜你做了什么判断｜输出是什么｜带来了什么实际变化。",
            lines=3,
        )


def add_contact(doc):
    add_section_intro(doc, "06", "Contact — Let’s Connect", "目标：给招聘者一个自然、明确、低阻力的下一步。")
    doc.add_heading("建议结尾文案", level=2)
    add_body(doc, "Let’s build something people want to join.")
    add_body(doc, "Have an idea, question, or interesting problem? Let’s connect.", color=MUTED)
    add_status_callout(
        doc,
        "现有信息：Email yanzhuofan333@163.com；Phone +86 139 0399 2996；Location Shanghai, China。请确认是否仍准确并允许公开。"
    )
    add_prefill(
        doc,
        "简历确认：手机（微信）13903992996；邮箱 yanzhuofan333@163.com。简历未提供 LinkedIn URL、个人微信二维码和英文简历 PDF。"
    )
    add_prompt(
        doc,
        "Email",
        "请确认最终邮箱；点击后将打开默认邮件客户端。",
        lines=1,
    )
    add_prompt(
        doc,
        "Phone",
        "请确认号码格式，以及是否公开显示。若不公开可删除。",
        lines=1,
    )
    add_prompt(
        doc,
        "WeChat",
        "请确认按钮显示文案，以及二维码是否长期有效。",
        lines=1,
    )
    add_prompt(
        doc,
        "LinkedIn 与简历",
        "请提供 LinkedIn URL；说明是否提供 Resume 下载，并上传最终 PDF。",
        lines=2,
    )
    add_prompt(
        doc,
        "联系区最后一句",
        "可保留建议文案，或写更像你本人语气的一句话。",
        lines=2,
    )
    add_photo_request(
        doc,
        "CONTACT-01",
        "微信二维码",
        "PNG/JPG 原图；二维码四周需有留白，确保手机扫码可识别。请确认不是短期群二维码。",
    )


def add_inventory(doc):
    add_section_intro(doc, "07", "照片与文件上传清单", "上传时请直接使用下列编号命名，后续放入网站会快很多。")
    rows = [
        ("HERO-01", "首页个人照片", "必需", "2–3 张候选"),
        ("ABOUT-01", "About Me 辅助照片", "可选", "1–2 张"),
        ("EXP-01", "Journey 目录/路线视觉", "必需", "1 张"),
        ("EXP-02", "本科", "必需", "1–3 张"),
        ("EXP-03", "实习一", "必需", "1–3 张"),
        ("EXP-04", "实习二", "必需", "1–3 张"),
        ("EXP-05", "研究生 / Boston", "必需", "1–3 张"),
        ("EXP-06", "工作 / Product & GTM", "必需", "1–3 张"),
        ("CASE01-01/02", "ZURU × Walmart", "必需", "4–7 张"),
        ("CASE02-01/02", "PepsiCo / Quaker", "必需", "4–7 张"),
        ("CASE04-01", "Journaling Community", "可选", "3–6 张"),
        ("AI-01/02", "Five-Year Diary", "必需", "6–10 张"),
        ("CONTACT-01", "微信二维码", "必需", "1 张"),
        ("RESUME-01", "英文简历 PDF", "建议", "1 个 PDF"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [1700, 4050, 1300, 2310], 120)
    set_table_borders(table)
    headers = ["编号", "用途", "优先级", "建议数量"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for row_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            if row_idx % 2 == 1:
                set_cell_shading(cells[idx], LIGHT_GRAY)
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=9)
            set_cell_margins(cells[idx], top=90, bottom=90)
    doc.add_paragraph()
    doc.add_heading("照片上传规则", level=2)
    for item in [
        "保留原图，不要先压缩；如需修图，只做基础亮度、裁切和色彩调整。",
        "不要使用带水印、社交平台 UI 或明显滤镜的截图作为主视觉。",
        "产品/公司项目素材必须确认可公开；敏感数据、客户名和内部界面请先脱敏。",
        "同一段经历尽量提供人物、环境和细节三种照片，方便网页排版选择。",
        "建议文件名：编号_地点或项目_年份_序号，例如 EXP-05_Boston_2022_01.jpg。",
    ]:
        add_bullet(doc, item)
    add_prompt(
        doc,
        "已有照片与编号对应关系",
        "如果你已经知道某张照片要放在哪里，请在这里写“原文件名 → 编号”。",
        lines=7,
    )


def add_final_check(doc):
    add_section_intro(doc, "08", "提交前最终确认", "不需要一次把所有英文写完；先保证信息真实、完整、可公开。")
    checks = [
        "Hero 主文案、姓名、城市和求职状态已确认",
        "About Me 已写清 PM × GTM 的差异化",
        "Experience 六页的时间、地点、组织和故事均已补齐",
        "快手、百事两段实习的公司官方英文名和个人贡献已确认",
        "Cases 的所有数据均有口径、时间范围和公开许可",
        "每个 Case 已区分个人贡献与团队贡献",
        "每个 Case 均回答 How I’d do it better now",
        "Five-Year Diary 链接、截图和产品机制已补齐",
        "AI 工作流使用的是真实例子，而非工具清单",
        "Email、电话、LinkedIn、微信二维码和简历已确认",
        "所有必需照片已经按编号上传",
    ]
    for item in checks:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run("☐ " + item)
        set_run_font(r, size=10.5)
    add_prompt(
        doc,
        "还有什么一定希望网站保留？",
        "可以是一个故事、一张照片、一句个人表达或一个项目。",
        lines=4,
    )
    add_prompt(
        doc,
        "有什么一定不希望网站呈现？",
        "例如不公开的信息、不喜欢的职业标签、希望避免的视觉风格。",
        lines=4,
    )


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_footer(doc.sections[0])

    add_cover(doc)
    add_page_break(doc)
    add_quick_start(doc)
    add_page_break(doc)
    add_hero(doc)
    add_page_break(doc)
    add_about(doc)
    add_page_break(doc)
    add_experience(doc)
    add_page_break(doc)
    add_cases(doc)
    add_page_break(doc)
    add_ai(doc)
    add_contact(doc)
    add_page_break(doc)
    add_inventory(doc)
    add_page_break(doc)
    add_final_check(doc)

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.78)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)
        section.header_distance = Inches(0.38)
        section.footer_distance = Inches(0.38)
        add_footer(section)

    doc.core_properties.title = "个人简历网站 内容与照片补充工作簿"
    doc.core_properties.subject = "Portfolio website content and photo collection workbook"
    doc.core_properties.author = "Zhuofan Yan"
    doc.core_properties.keywords = "portfolio, website, content, photos, experience, case study"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
